from tkinter import *
from tkinter import ttk
from tkinter import messagebox
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys

global doc
import json

doc = docx.Document("template.docx")


def k():
    para2 = doc.paragraphs[0]
    t = para2.text
    t = t.replace("教科", combobox.get())
    para2.text = t

    print("c1")


def title():
    # タイトルの文字の大きさを変更する必要あり
    global table
    table = doc.tables[0]

    def replace_table_cell_text(cell, old_text, new_text):
        cell.text = cell.text.replace(old_text, new_text)

    replace_table_cell_text(table.cell(0, 0), "タイトル", TitleE.get())
    tbl = doc.tables[0]
    header = tbl.rows[0]  # 表の1行目
    for cell in header.cells:
        # セルの中の段落
        cell_para = cell.paragraphs[0]
        for run in cell_para.runs:
            if run.text:
                run.font.size = docx.shared.Pt(24)
                run.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # para5 = doc.paragraphs[13]
    # t5 = para5.text
    # t5 = t5.replace("タイトル", TitleE.get())
    # para5.text = t5
    # print("c2")


def name():
    global namaeE
    para3 = doc.paragraphs[21]
    t1 = para3.text
    t1 = t1.replace("名前", namaeE.get())
    para3.text = t1
    print("c3")


def number():
    global NumberE
    para4 = doc.paragraphs[20]
    t2 = para4.text
    t2 = t2.replace("ナンバー", NumberE.get())
    para4.text = t2
    print("c4")


def resize():
    doc.paragraphs[0].runs[0].font.size = docx.shared.Pt(11)
    doc.paragraphs[20].runs[0].font.size = docx.shared.Pt(14)
    doc.paragraphs[21].runs[0].font.size = docx.shared.Pt(14)
    # doc.tables[0].cell[0].runs[0].font.size = docx.shared.Pt(20)

    print("c5")


def save():
    doc.save(combobox.get() + ".docx")


def show_selected(event):
    print(combobox.get())


# =========================================================
# =========================================================
def clicked():
    k()
    name()
    number()
    resize()
    title()
    save()


root = Tk()
root.title("マイプロテンプレート作成")
root.geometry("400x300")

# オブジェクトの定義
namaeL = ttk.Label(root, text="名前を入力")
namaeE = ttk.Entry(root)
TitleL = ttk.Label(root, text="タイトルを入力")
TitleE = ttk.Entry(root)
NumberL = ttk.Label(root, text="学籍番号を入力")
NumberE = ttk.Entry(root)
Button = ttk.Button(
    root,
    text="ok",
    command=clicked,
)


item_list = [
    "古文",
    "数学",
    "地理",
    "英語",
]
KyoukaL = ttk.Label(root, text="教科を選択")
combobox = ttk.Combobox(master=root, values=item_list)
combobox.bind(
    "<<ComboboxSelected>>",
    show_selected,
)

namaeE.bind()

namaeL.pack()
namaeE.pack()
NumberL.pack()
NumberE.pack()
KyoukaL.pack()
combobox.pack()
TitleL.pack()
TitleE.pack()
Button.pack()
# ウィンドウの表
root.mainloop()
